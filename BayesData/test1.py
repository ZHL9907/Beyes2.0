import math
import cv2
import numpy as np
from matplotlib import pyplot as plt
import tkinter as tk
from tkinter import scrolledtext as st
import tkinter.filedialog as file
import tkinter.messagebox as mes
import jieba
import numpy as np
import docx
from tkinter import scrolledtext as st

def OpenDoc():
    global words
    #1.弹出打开文件框
    filename = file.askopenfilename(title='打开文件名字', initialdir="J:\pyproject",
                                    filetypes=[('docx文件', '*.docx')])
    #2.返回一个document对象doc1，返回的是一段段的
    doc1=docx.Document(filename)
    fullText=[]

    #3.将该文件得每一个句子拼接起来，成为一个fullText
    #doc1.paragraphs可以识别文档内得句子
    for para in doc1.paragraphs:
        fullText.append(para.text)
    print("len",len(fullText))

    #4.将fullText变成数组ft
    ft=np.array(fullText)
    cft=''
    bd = ',!#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~]+。，！？“”《》：、． '

    #5.将数组ft变成cft字符串
    for i in range(len(ft)):
         cft=cft+ft[i]

    #6.将字符串cft中得特殊字符消除,去除标点符号，变成纯文字
    for i in bd:
        cft = cft.replace(i,'')

    #7.将ctf进行jieba分词后生成word
    words= jieba.lcut(cft)

    #8.清空ScrolledText（滚动条） std并将word插入
#    std.delete(1.0,tk.END)
#    std.insert(tk.END, words)
    print(words)

    return

#添加一个训练文档
def AddDoc():
    global TagFlag,TrainDoc,DocTag,words,TargetVector,CountStr
    if (len(words) == 0 ):
        return
    #选择了Radiobutton后TagFlag会变成true
    if(TagFlag==False):
        mes.askokcancel(title="警告",message="没有选择类别")
        return
    #1.将words添加到训练样本对象TrainDoc中
    #TrainDoc：全局变量，样本
    TrainDoc.append(words)

    #TargetVector是一个十维0向量，当选择某一个标签时，该标签对应得向量位置置为1
    #2.将向量拼接到DocTag中
    DocTag.append(TargetVector)

    #3.获取训练样本得词数？
    #CountStr全局变量，样本Doc数
    CountStr.set(str(len(TrainDoc)))

    TagFlag = False
    return


OpenDoc()