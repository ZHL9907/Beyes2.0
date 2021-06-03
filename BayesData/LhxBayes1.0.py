#!/usr/bin/env python
# -*- coding:utf-8 -*-
#This Python codes are created by Li hexi for graduate student course --Artificial Intelligence and Machine learning
#Any copy of the codes is illegal without the author' permission.2020-4-3
#--------This program uses naive Bayes principle to classify documents------
import math
import cv2
import numpy as np
from matplotlib import pyplot as plt
import tkinter as tk
from tkinter import scrolledtext as st
import tkinter.filedialog as file
import tkinter.messagebox as mes
import jieba
import numpy as np
import docx
from tkinter import scrolledtext as st
root = tk.Tk()
ww=600;
wh=500;
TargetVector=[]
TagFlag=False
TrainFlag=False
FinishFlag=False
CountStr=tk.StringVar()
DigitStr=tk.StringVar()
EpochStr=tk.StringVar()
ErrStr=tk.StringVar()
root.title('机器学习')
root.geometry('1200x800')
root.resizable(width=False,height=False)
frma1 = tk.Frame(root, width=ww+50, height=wh+50, relief=tk.GROOVE)
frma1.place(relx=0, rely=0)

I1=cv2.imread('image/Bayeslogo.png')
I2=cv2.resize(I1,(ww,wh))
cv2.imwrite('image/logo.png', I2)
image_file = tk.PhotoImage(file='image/logo.png')
L1=tk.Label(frma1, image=image_file)
L1.place(x=50, y=50)
std=st.ScrolledText(root,width=25,height=15,font=("宋体",20))
std.place(x=700,y=90)
#doc=docx.Document()
ClassNum=10#文档分类的类别数
TrainDoc=[]#训练文档集
DocTag=[]#训练文档的标签


"""
1.打开样本
2.打标签
3.添加样本
4.贝叶斯学习
5.贝叶斯分类
"""
words=[]#临时打开的文档
#打开一个中文word文档作为样本
def OpenDoc():
    global words
    #1.弹出打开文件框
    filename = file.askopenfilename(title='打开文件名字', initialdir="J:\pyproject",
                                    filetypes=[('docx文件', '*.docx')])
    #2.返回一个document对象doc1
    doc1=docx.Document(filename)
    fullText=[]

    #3.将该文件得每一个句子拼接起来，成为一个fullText
    #doc1.paragraphs可以识别文档内得句子
    for para in doc1.paragraphs:
        fullText.append(para.text)
    print("len",len(fullText))      #len(fullText)因为上面是三段(数组)所以 len=3

    #4.将fullText变成数组ft
    ft=np.array(fullText)
    cft=''
    bd = ',!#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~]+。，！？“”《》：、． '

    #5.将数组ft变成cft字符串 [123124]→123124
    for i in range(len(ft)):
         cft=cft+ft[i]      #把三段连起来

    #6.将字符串cft中得特殊字符消除   我，你，它→我你它（变纯文本）
    for i in bd:
        cft = cft.replace(i,'')

    #7.将ctf进行jieba分词后生成word   jieba通过算法把，”我你它“ 分段
    words= jieba.lcut(cft)

    #8.清空ScrolledText（滚动框） std并将word插入
    std.delete(1.0,tk.END)
    std.insert(tk.END, words)
    return

#添加一个训练文档
def AddDoc():
    global TagFlag,TrainDoc,DocTag,words,TargetVector,CountStr
    if (len(words) == 0 ):
        return
    #选择了Radiobutton后TagFlag会变成true
    if(TagFlag==False):
        mes.askokcancel(title="警告",message="没有选择类别")
        return
    #1.将words添加到训练样本对象TrainDoc中
    #TrainDoc：全局变量，样本
    TrainDoc.append(words)

    #TargetVector是一个十维0向量，当选择某一个标签时，该标签对应得向量位置置为1
    #2.将向量拼接到DocTag中
    DocTag.append(TargetVector)  #TargetVector目标向量

    #3.获取训练样本得词数？
    #CountStr全局变量，样本Doc数
    CountStr.set(str(len(TrainDoc)))        #样本计数
    TagFlag = False
    return

#保存训练文档样本(训练模型)
def SaveDoc():
    global TrainDoc,DocTag
    if (len(TrainDoc)== 0):
        return
    file1 = file.asksaveasfilename()
    if (file1 == ""):
        return
    np.savez(file1+".npz", Doc=TrainDoc,Tag=DocTag)


#有不懂
#装入入训练文档样本
def LoadDoc( ):
    global TrainDoc, DocTag
    file1 = file.askopenfilename(title='打开文件名字', initialdir="\pyproject\Bayes",
                                 filetypes=[('npz文件', '*.npz')])
    if(file1==""):
        return
    ls=np.load(file1,allow_pickle=True) #allow_pickle允许多文件，因为保存npz时有多文件
    TrainDoc=ls['Doc']  #???
    #TrainDoc=list(ls['Doc']) 把保存的文件转换成python的列表格式
    DocTag=ls['Tag']    #???
    CountStr.set(str(len(TrainDoc)))    #训练集里有多少样本
    std.delete(1.0, tk.END)
    std.insert(tk.END,TrainDoc)

#贝叶斯学习 （算概率）
def BayesLearn( ):
    global TrainDoc, DocTag, TrainFlag,ClassNum
    if (len(TrainDoc)==0):
        return

    #1.DocTag中，每有一个样本，就有一个向量，获取样本得数量
    DocNum=len(DocTag)#文档总数
    ClassNum = 10# 分类数
    n=np.zeros(ClassNum) #类别初始值为0，其实建立10个【0】数组
    global pv
    pv=np.zeros(ClassNum)
    lt=[]

    #2.第一层，遍历每个训练样本
    #  第二层，遍历训练样本中的每一个词，将每一个词加入到lt
    m=len(TrainDoc) #文档计数
    for i in range(m):
        for j in range(len(TrainDoc[i])):   #文档里字长。
            lt.append(TrainDoc[i][j])       #添加到lt=[]里，i代表第几个文档，j代表第几个单词，一个个单词连起来
                                            #把所有文档放到lt[]列表中。
    #3.建立词库，消除重复单词
    global wordlib
    wordlib=np.unique(lt)#建立词库，消除重复单词
    WordNum=len(wordlib)#词库单词数|Vocabulary|


    #4.计算同类文档的先验概率pv
    #第一层循环：遍历文档数
    #第二层循环：遍历文档的类型（0，1，2）
    TextClass=[[] for row in range(ClassNum)]#每个分类文档的联合列表的初始状态,二维列表空表
    for i in range(DocNum):#所有文档循环
        for j in range(ClassNum):#把类别相同的文档串一起
            if DocTag[i][j]==1.0:#训练文档的类别标签，属于哪个类别就是==1.0
                TextClass[j].append(TrainDoc[i])#将同类文档连接起来，注意这是一个三维list
                n[j]=n[j]+1 #统计同类文档个数

    #5.变量n[j]：同类文档i的数目
    for i in range(ClassNum):
        pv[i]=n[i]/(n.sum())#相当讲义的P(vj)
    print(pv)

    #计算类条件概率pk
    global pk
    pk=np.zeros((WordNum,ClassNum))#相当讲义的P(wk|vj)

    #第一层循环：遍历所有词
    #第二层循环，遍历类型
    #第三层循环：遍历每一种类型的文档数
    #第四层循环：遍历每一种类型的文档的每一个单词
    for i in range(WordNum):
         for j in range(ClassNum):
             nk=0#在某类文档中出现的次数
             nn=0#在某类文档单词总数
             for k1 in range(len(TextClass[j])):
                 m=len(TextClass[j][k1])
                 nn=nn+m # nn文档中第J类单词总计数
                 for k2 in range (m):
                     if wordlib[i]==TextClass[j][k1][k2]:  #10个类型，每个类型其中一个文件，一个文件内第几个词
                           nk=nk+1# nk在J类文档中单词wk 出现的次数
             #nn文档中第J类单词总计数
             #nk在J类文档中单词wk 出现的次数
             #P(wk|vj)
             pk[i][j]=(nk+1.0)/(nn+WordNum)#采纳m-估计方法

    np.savez("pvpk.npz",pv=pv,pk=pk,wordlib=wordlib)#将先验概率、条件概率和词库一起暂时存盘，pv,pk,wordlib包装成npz文件

    #读取pvpk.npz，并打开对应的Doc文件方法，无需去注释
    #ls=np.load(file1,allow_pickle=True)
    #TrainDoc=ls['Doc']

    TrainFlag=True #训练完成标志
    print(pk)
    return

def SaveTrainResult():
    global pv,pk,wordlib
    if TrainFlag==False:
        return
    file1 = file.asksaveasfilename()
    if (file1 == ""):
        return
    np.savez(file1+".npz", pv=pv, pk=pk, wordlib=wordlib)  # 将先验概率、条件概率和词库一起存起个名字存盘
    return
def LoadTrainResult():
    global pv, pk, wordlib,TrainFlag
    file1 = file.askopenfilename(title='打开文件名字', initialdir="\pyproject\Bayes",
                                 filetypes=[('npz文件', '*.npz')])
    if (file1 == ""):
        TrainFlag = False
        return
    ls = np.load(file1, allow_pickle=True)
    pv = ls['pv']
    pk = ls['pk']
    wordlib = ls['wordlib']
    print(pk)
    TrainFlag=True
    return

def BayesClassify( ):
    global TrainFlag,words,DigitStr,pv,pk,wordlib
    #print(pk)
    if(TrainFlag==False):
        mes.askokcancel(title="警告", message="没有训练") 
        return
    if(len(words)==0):
        mes.askokcancel(title="警告", message="没有文档") 
        return
    #ls = np.load("pvpk.npz", allow_pickle=True)
    #pv = ls['pv']
    #pk = ls['pk']
    #wordlib=ls['wordlib']

    ClassNum=len(pv)
    WordNum=np.size(pk,axis=0)
    p=np.ones(ClassNum)

    #第一层循环：遍历所有词(jieba处理过的词)(新打开的文本)
    #第二层循环：找单词在词汇表中的位置（已经学习好的训练集）
    #第三层循环：计算所有本单词属于的类别条件概率
    for i in range(len(words)):#按position位置(属性）循环
        for j in range(WordNum):#寻找单词在词汇表中的位置
            if words[i]==wordlib[j]:
               for k in range(ClassNum):#计算所有本单词属于的类别条件概率
                   p[k]=p[k]*pk[j][k]
                   while p[k]<0.000001:#防止连乘结果过小比例放大pk
                       p=p*10.0

    for i in range(ClassNum):
        p[i]=pv[i]*p[i]
    print(p)
    #numpy.argmax(array, axis) 用于返回一个numpy数组中最大值的索引值。当一组中同时出现几个最大值时，返回第一个最大值的索引值
    imax = np.argmax(p,axis=0)
    print(imax)
    Result = ['政治', '经济','军事','教育','科技','体育','艺术','旅游','汽车','娱乐']
    DigitStr.set(Result[imax])
    root.mainloop()
    return

def CallBack():
    return
def picshow(filename):
    global GI
    I1 = cv2.imread(filename)
    I2 = cv2.resize(I1,(ww, wh))
    cv2.imwrite('image/temp.png', I2)
    I3 = tk.PhotoImage(file='image/temp.png')
    L1 = tk.Label(frma1, image=I3)
    L1.place(x=50, y=50)
    GI=np.array(I2)
    n1,n2,n3=np.shape(GI)
    print(n1,n2,n3)
    root.mainloop()
menubar=tk.Menu(root)
#file menu
fmenu=tk.Menu(menubar)
fmenu.add_command(label = '新建',command=CallBack)
fmenu.add_command(label = '打开',command=CallBack)
fmenu.add_command(label = '保存',command=CallBack)
fmenu.add_command(label = '另存为',command=CallBack)
#Image processing menu
imenu=tk.Menu(menubar)
imenu.add_command(label='FFT变换',command=CallBack)
imenu.add_command(label='DOT变换',command=CallBack)
imenu.add_command(label='边缘检测',command=CallBack)
imenu.add_command(label='区域分割',command=CallBack)
#machine learning
mmenu=tk.Menu(menubar)
mmenu.add_command(label = 'KNN',command=CallBack)
mmenu.add_command(label = '朴素贝叶斯',command=CallBack)
mmenu.add_command(label = '支持向量机',command=CallBack)
mmenu.add_command(label = 'BP神经网',command=CallBack)
mmenu.add_command(label = 'CNN卷积神经网',command=CallBack)
bmenu=tk.Menu(menubar)
bmenu.add_command(label = '装入文档',command=OpenDoc)
bmenu.add_command(label = '添加文档样本',command=AddDoc)
bmenu.add_command(label = '保存文档样本',command=SaveDoc)
bmenu.add_command(label = '装入文档样本',command=LoadDoc)
bmenu.add_command(label = '贝叶斯学习',command=BayesLearn)
bmenu.add_command(label = '贝叶斯分类',command=BayesClassify)
menubar.add_cascade(label = "文件操作",menu = fmenu)
menubar.add_cascade(label = "图像处理",menu = imenu)
menubar.add_cascade(label = "机器学习",menu = mmenu)
menubar.add_cascade(label = "朴素贝叶斯学习",menu = bmenu)
root.config(menu=menubar)

but1=tk.Button(root, text='打开文档', font=('Arial', 12), width=10, height=1, command=OpenDoc)
but2=tk.Button(root, text='添加样本', font=('Arial', 12), width=10, height=1, command=AddDoc)
but3=tk.Button(root, text='保存样本', font=('Arial', 12), width=10, height=1, command=SaveDoc)
but4=tk.Button(root, text='装入样本', font=('Arial', 12), width=10, height=1, command=LoadDoc)
but5=tk.Button(root, text='Bayes学习', font=('Arial', 12), width=10, height=1, command=BayesLearn)
but6=tk.Button(root, text='Bayes分类', font=('Arial', 12), width=10, height=1, command=BayesClassify)
but7=tk.Button(root, text='保存训练结果', font=('Arial', 12), width=10, height=1, command=SaveTrainResult)
but8=tk.Button(root, text='装入训练结果', font=('Arial', 12), width=10, height=1, command=LoadTrainResult)
but9=tk.Button(root, text='备用按钮', font=('Arial', 12), width=10, height=1, command=CallBack)
dd=40
but1.place(x=100,y=560+dd)
but2.place(x=300,y=560+dd)
but3.place(x=500,y=560+dd)
but4.place(x=100,y=620+dd)
but5.place(x=300,y=620+dd)
but6.place(x=500,y=620+dd)
but7.place(x=100,y=680+dd)
but8.place(x=300,y=680+dd)
but9.place(x=500,y=680+dd)
#Text1=tk.Text(root)
#Text1.place(x=850,y=500)
entry1=tk.Entry(root,font=('Arial', 12),width=15, textvariable=CountStr)
entry1.place(x=900,y=600)
CountStr.set("0")
entry2=tk.Entry(root,font=('Arial', 12), width=15,textvariable=EpochStr)
entry2.place(x=900,y=660)
EpochStr.set("0")
entry3=tk.Entry(root,font=('Arial', 12), width=15,textvariable=ErrStr)
entry3.place(x=900,y=720)
ErrStr.set("0")
entry4=tk.Entry(root,font=('Arial', 12), width=10,textvariable=DigitStr)
entry4.place(x=260,y=10)
DigitStr.set('')
Lab1=tk.Label(root,text='样本计数:',font=('Arial', 12),width=10, height=1)
Lab1.place(x=780, y=600)
Lab2=tk.Label(root,text='训练次数:',font=('Arial', 12),width=10, height=1)
Lab2.place(x=780, y=660)
Lab3=tk.Label(root,text='训练误差:',font=('Arial', 12),width=10, height=1)
Lab3.place(x=780, y=720)
Lab4=tk.Label(root,text='识别结果:',font=('Arial', 12),width=10, height=1)
Lab4.place(x=100, y=10)
Lab5=tk.Label(root,text='文档显示窗口',font=('Arial', 14),width=10, height=1)
Lab5.place(x=850, y=50)
#text1=tk.Text(root, font=('Arial', 12),fg='black', width=10, height=1)
#text1.place(x=800,y=400)
#设置目标标签无线按钮
Target=[('政治',0),('经济',1),('军事',2),('教育',3),('科技',4),('体育',5),('艺术',6),('旅游',7),('汽车',8),('娱乐',9)]
Target_startx=5
Target_starty=560
TargetVector=np.zeros(10)
TargetV=tk.IntVar()

#标签判断
def SetTarget():
    global TagFlag,TargetVector
    TagFlag=True
    for i in range(10):
        if (TargetV.get()==i):
            TargetVector=np.zeros(10) #其他未选择的为0，选中RadioButtion为1.0
            TargetVector[i]=1.0
            print(TargetVector)

for txt,num in Target:
    rbut=tk.Radiobutton(root, text=txt, value=num,font=('Arial', 12), width=3, height=1,command=SetTarget, variable=TargetV)
    rbut.place(x=Target_startx+num*70,y=Target_starty)
#root.after(1000,PlotLine)
root.mainloop()